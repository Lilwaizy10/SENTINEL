from docx import Document
import sys

doc = Document(r'c:\Users\qaizu\Documents\DLW Hackathon\SENTINEL_Build_Spec (1).docx')

# Extract paragraphs
for p in doc.paragraphs:
    text = p.text
    if text:
        sys.stdout.buffer.write((text + '\n').encode('utf-8', errors='replace'))

# Extract tables
for table in doc.tables:
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            if cell.text:
                row_data.append(cell.text)
        if row_data:
            sys.stdout.buffer.write((' | '.join(row_data) + '\n').encode('utf-8', errors='replace'))
